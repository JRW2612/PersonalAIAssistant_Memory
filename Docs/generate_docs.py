import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_styled_doc(title, subtitle):
    doc = docx.Document()
    
    # Page Setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Segoe UI'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x2D, 0x37, 0x48) # Slate dark
    
    # Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run(title)
    run_title.font.name = 'Segoe UI'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Deep Navy

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run(subtitle)
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    
    return doc

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0) # Blue
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
    return h

def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.bold = True
        r_bold.font.color.rgb = RGBColor(0x1A, 0x20, 0x2C)
    p.add_run(text)
    return p

def add_styled_table(doc, headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header Row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        set_cell_background(hdr_cells[i], "1A365D")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.name = 'Segoe UI'
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data Rows
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_value in enumerate(row_data):
            row_cells[c_idx].text = str(cell_value)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

out_dir = r"d:\MyProjects\PersonalAIAssistant_Memory\Docs"
os.makedirs(out_dir, exist_ok=True)

# ---------------------------------------------------------
# 1. SYSTEM DESIGN DOCUMENT
# ---------------------------------------------------------
doc1 = create_styled_doc(
    "System Design Document", 
    "Personal AI Assistant — Long-Term Memory Core Engine v2.0"
)

add_heading_1(doc1, "1. Executive Architecture Overview")
add_paragraph(doc1, "The Personal AI Assistant Memory Core is an event-sourced, high-throughput, multi-tenant memory management microservice. Designed to solve conversational AI amnesia ('Goldfish Syndrome'), the system provides persistent, context-aware memory retrieval and background consolidation for autonomous AI agents.")

add_heading_1(doc1, "2. Domain-Driven CQRS & Event Sourcing Architecture")
add_paragraph(doc1, "The core system strictly decouples Write Operations (Commands/Events) from Read Operations (Queries/Vector Retrieval) using the Command Query Responsibility Segregation (CQRS) pattern.")
add_bullet(doc1, "Stores immutable domain events (MemoryAdded, MemoryCompressed, MemoryConsolidated, MemoryArchived) in MongoDB. Serves as the single source of truth for full aggregate history playback.", "Event Store (MongoDB): ")
add_bullet(doc1, "Maintains relational projections of memory metadata in PostgreSQL for high-speed indexing, count queries, and batch candidates.", "Read Model (PostgreSQL): ")
add_bullet(doc1, "Stores 1536-dimensional vector embeddings with Int8 Scalar Quantization for low-latency semantic similarity search.", "Vector Store (Qdrant): ")
add_bullet(doc1, "Asynchronous event distribution over MassTransit & RabbitMQ, ensuring eventual consistency between Write and Read models.", "Message Broker (RabbitMQ): ")

add_heading_1(doc1, "3. Security & Multi-Tenancy Design")
add_bullet(doc1, "All event payloads stored in MongoDB and summary texts stored in PostgreSQL are encrypted using AES-256-CBC with per-user derived keys (SystemKey + UserId).", "Transparent Data Encryption (TDE): ")
add_bullet(doc1, "All Qdrant vector similarity queries strictly enforce a mandatory payload filter (Key: 'userId') to prevent cross-tenant memory leakage.", "Strict Vector Isolation: ")
add_bullet(doc1, "HTTP Clients (OpenAI, Gemini, Teams Webhook) are wrapped with Polly Resilience Handlers providing exponential backoff retries, rate limiting, and circuit breaker protection.", "Resilience Pipelines: ")

add_heading_1(doc1, "4. Technology Stack & Component Specifications")
headers1 = ["Layer / Component", "Technology Selected", "Purpose & Architectural Role"]
data1 = [
    ["Runtime Framework", ".NET 8.0 (LTS)", "Core high-performance backend microservice"],
    ["Event Store", "MongoDB 3.9", "Immutable append-only domain event storage"],
    ["Read Repository", "PostgreSQL 16 / EF Core 8", "CQRS read model projections & fast query indexing"],
    ["Vector Engine", "Qdrant gRPC (Client v1.19)", "Semantic similarity retrieval & Quantized embedding storage"],
    ["Message Broker", "MassTransit + RabbitMQ 9.2", "Decoupled domain event broadcasting & asynchronous processing"],
    ["Resilience & Fault Tolerance", "Polly 8.8 / Microsoft.Extensions.Http", "Circuit breaker, rate limiting, and backoff retries for external AI APIs"]
]
add_styled_table(doc1, headers1, data1)

doc1.save(os.path.join(out_dir, "System_Design_Document.docx"))

# ---------------------------------------------------------
# 2. BUSINESS RULES SPECIFICATION
# ---------------------------------------------------------
doc2 = create_styled_doc(
    "Business Rules & Domain Specification", 
    "Personal AI Assistant — Memory Governance, Lifecycle & Policy Engine"
)

add_heading_1(doc2, "1. Memory Ingestion & Chunking Directives")
add_bullet(doc2, "When a memory text exceeds configured token limits (default: 500 tokens), the TextChunker engine splits the raw input using a sliding window overlap (default: 50 tokens overlap).", "BR-01 Sliding Window Chunking: ")
add_bullet(doc2, "Each chunk is instantiated as an independent MemoryAggregate with a unique Stream ID (memory-{Guid}), tagged with 'chunk:N' and 'parent:{CorrelationId}' for upstream tracing.", "BR-02 Correlation & Lineage: ")

add_heading_1(doc2, "2. Memory Lifecycle & Consolidation Policies")
add_bullet(doc2, "The ConsolidationWorker background job queries candidate memories with >50 tokens and processes them via LLM compression to synthesize dense, high-utility summaries.", "BR-03 Background Consolidation: ")
add_bullet(doc2, "Background workers use IServiceScopeFactory per iteration to prevent captive dependency bugs and memory leaks.", "BR-04 Scoped Worker Isolation: ")
add_bullet(doc2, "RetentionWorker enforces a 30-day Time-To-Live (TTL) on standard tier memories, automatically flagging expired items for archiving.", "BR-05 TTL & Expiration: ")

add_heading_1(doc2, "3. Multi-Tenant Capacity & Subscription Tiers")
headers2 = ["Subscription Tier", "Max Memories Capacity", "Retention Policy (TTL)", "LLM Compression & Consolidation"]
data2 = [
    ["Free / Community Tier", "100 Active Memories", "30 Days TTL Expiration", "Standard Deterministic Compression"],
    ["Professional Tier", "5,000 Active Memories", "365 Days Retention", "Background LLM Consolidation"],
    ["Enterprise Tier", "Unlimited Memories", "Lifetime Infinite Retention", "Advanced Semantic Graph & Vector Synthesis"]
]
add_styled_table(doc2, headers2, data2)

doc2.save(os.path.join(out_dir, "Business_Rules_Specification.docx"))

# ---------------------------------------------------------
# 3. COST ESTIMATION & FINOPS REPORT
# ---------------------------------------------------------
doc3 = create_styled_doc(
    "Cloud Infrastructure & FinOps Cost Estimation", 
    "Personal AI Assistant — Scalability & Cost Modeling (10k to 1M DAU)"
)

add_heading_1(doc3, "1. FinOps Optimization Strategies Implemented")
add_bullet(doc3, "By configuring Qdrant Int8 Scalar Quantization, vector precision is converted from float32 to int8, reducing RAM consumption by 75% with negligible similarity accuracy loss.", "Qdrant Vector RAM Reduction: ")
add_bullet(doc3, "ConsolidationWorker compresses 50+ memory items into 1 dense summary, reducing downstream retrieval prompt context tokens by up to 80%.", "LLM Token Savings: ")

add_heading_1(doc3, "2. Estimated Infrastructure Budget Scaling")
headers3 = ["Active Daily Users (DAU)", "Vector DB RAM (Qdrant)", "Primary DB (Mongo + Postgres)", "LLM Token API Cost (Monthly)", "Total Monthly Budget"]
data3 = [
    ["10,000 Users", "4 GB RAM ($40/mo)", "2 Nodes Shared ($60/mo)", "$150 / month", "~ $250 / month"],
    ["100,000 Users", "32 GB RAM ($280/mo)", "Managed Cluster ($350/mo)", "$1,200 / month", "~ $1,830 / month"],
    ["1,000,000 Users", "256 GB RAM ($1,900/mo)", "Enterprise Sharded Cluster ($2,500/mo)", "$9,500 / month", "~ $13,900 / month"]
]
add_styled_table(doc3, headers3, data3)

doc3.save(os.path.join(out_dir, "Cost_Estimation_FinOps_Report.docx"))

# ---------------------------------------------------------
# 4. PRE-PROJECT RESEARCH REPORT
# ---------------------------------------------------------
doc4 = create_styled_doc(
    "Pre-Project Research & State-of-the-Art Analysis", 
    "Personal AI Assistant — Solving Conversational Memory Amnesia in AI Systems"
)

add_heading_1(doc4, "1. Executive Problem Statement: Goldfish Syndrome")
add_paragraph(doc4, "Traditional Large Language Models (LLMs) operate with stateless inference windows. Standard context windows suffer from high token consumption, high latency, context degradation ('Lost in the Middle'), and complete lack of persistent user memory across sessions.")

add_heading_1(doc4, "2. Paradigmatic Comparison")
headers4 = ["Dimension", "Generative AI (Stateless)", "Agentic AI (Task-Driven)", "AI Agents with Persistent Memory (Our Architecture)"]
data4 = [
    ["Context Window", "Single Prompt Session", "Workflow-bound Execution", "Infinite Historical Continuum"],
    ["State Management", "Stateless / Transient", "Short-Term Working Memory", "Event-Sourced Long-Term Memory Core"],
    ["Multi-Session Learning", "None (Amnesia)", "Task-scoped", "Adaptive User-Specific Learning"],
    ["Retrieval Model", "Direct Prompt Injection", "Tool Calling / Ephemeral RAG", "Dynamic Hybrid RAG (Vector + Temporal + Importance)"]
]
add_styled_table(doc4, headers4, data4)

add_heading_1(doc4, "3. Key Industry Benchmarks & Research Citations")
add_bullet(doc4, "Examines why AI models suffer from amnesia and highlights the necessity of decoupled read/write stores for long-term agent memory.", "Oracle Developers Research: ")
add_bullet(doc4, "Highlights community demand for persistent user context across OpenAI ChatGPT sessions.", "OpenAI Forum Benchmarks: ")
add_bullet(doc4, "Systematic analysis of memory architectures, recommending hybrid event-sourced vector pipelines.", "Jason Scott Montoya Memory Analysis: ")

doc4.save(os.path.join(out_dir, "Pre_Project_Research_Report.docx"))

print("All 4 DOCX files successfully created in:", out_dir)
